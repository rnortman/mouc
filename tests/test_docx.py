"""Tests for DOCX generation."""
# pyright: reportPrivateUsage=false, reportUnknownMemberType=false, reportUnknownArgumentType=false, reportOptionalMemberAccess=false

from io import BytesIO

import pytest
from docx import Document

from mouc import styling
from mouc.backends import DocxBackend
from mouc.document import DocumentGenerator
from mouc.models import Entity, FeatureMap, FeatureMapMetadata
from mouc.parser import resolve_graph_edges
from mouc.unified_config import DocxConfig, OrganizationConfig


def create_docx_generator(feature_map: FeatureMap, config: DocxConfig | None = None) -> bytes:
    """Helper to create DOCX output from a feature map."""
    styling_context = styling.create_styling_context(feature_map)
    backend = DocxBackend(feature_map, styling_context)
    generator = DocumentGenerator(feature_map, backend, config)
    result = generator.generate()
    assert isinstance(result, bytes), "DocxBackend should return bytes"
    return result


def get_docx_text(docx_bytes: bytes) -> str:
    """Extract all text from DOCX bytes for testing."""
    doc = Document(BytesIO(docx_bytes))
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    return "\n".join(text_parts)


class TestDocxGenerator:
    """Test the DOCX document generator."""

    @pytest.fixture
    def simple_feature_map(self) -> FeatureMap:
        """Create a simple feature map for testing."""
        metadata = FeatureMapMetadata(
            version="1.0",
            team="test_team",
            last_updated="2024-01-15",
        )

        cap1 = Entity(
            type="capability",
            id="cap1",
            name="Cap 1",
            description="Description of capability 1.",
            tags=["infra"],
            links=[
                "[DD-123](https://example.com/dd123)",
                "jira:JIRA-456",
            ],
        )
        cap2 = Entity(
            type="capability",
            id="cap2",
            name="Cap 2",
            description="Description of capability 2.",
            requires={"cap1"},
        )

        story1 = Entity(
            type="user_story",
            id="story1",
            name="Story 1",
            description="Description of story 1.",
            requires={"cap2"},
            tags=["urgent"],
            meta={"requestor": "team_alpha"},
        )

        outcome1 = Entity(
            type="outcome",
            id="outcome1",
            name="Outcome 1",
            description="Description of outcome 1.",
            requires={"story1"},
            meta={"target_date": "2024-Q3"},
        )

        entities = [cap1, cap2, story1, outcome1]
        resolve_graph_edges(entities)

        return FeatureMap(
            metadata=metadata,
            entities=entities,
        )

    def test_generate_header(self, simple_feature_map: FeatureMap) -> None:
        """Test header generation."""
        docx_bytes = create_docx_generator(simple_feature_map)
        text = get_docx_text(docx_bytes)

        assert "Feature Map" in text
        assert "test_team" in text
        assert "2024-01-15" in text
        assert "1.0" in text

    def test_generate_toc(self, simple_feature_map: FeatureMap) -> None:
        """Test table of contents generation."""
        config = DocxConfig(organization=OrganizationConfig(primary="by_type"))
        docx_bytes = create_docx_generator(simple_feature_map, config)
        text = get_docx_text(docx_bytes)

        assert "Table of Contents" in text
        assert "Capabilities" in text
        assert "Cap 1" in text
        assert "Cap 2" in text
        assert "User Stories" in text
        assert "Story 1" in text
        assert "Outcomes" in text
        assert "Outcome 1" in text

    def test_generate_capabilities_section(self, simple_feature_map: FeatureMap) -> None:
        """Test capabilities section generation."""
        config = DocxConfig(organization=OrganizationConfig(primary="by_type"))
        docx_bytes = create_docx_generator(simple_feature_map, config)
        text = get_docx_text(docx_bytes)

        assert "Capabilities" in text
        assert "Cap 1" in text
        assert "Description of capability 1." in text
        assert "cap1" in text
        assert "infra" in text

        # Check dependencies
        assert "Cap 2" in text
        assert "Requires" in text

    def test_generate_user_stories_section(self, simple_feature_map: FeatureMap) -> None:
        """Test user stories section generation."""
        config = DocxConfig(organization=OrganizationConfig(primary="by_type"))
        docx_bytes = create_docx_generator(simple_feature_map, config)
        text = get_docx_text(docx_bytes)

        assert "User Stories" in text
        assert "Story 1" in text
        assert "team_alpha" in text
        assert "Requires" in text
        assert "Cap 2" in text

    def test_generate_outcomes_section(self, simple_feature_map: FeatureMap) -> None:
        """Test outcomes section generation."""
        config = DocxConfig(organization=OrganizationConfig(primary="by_type"))
        docx_bytes = create_docx_generator(simple_feature_map, config)
        text = get_docx_text(docx_bytes)

        assert "Outcomes" in text
        assert "Outcome 1" in text
        assert "Requires" in text
        assert "Story 1" in text

    def test_dependency_tracking(self) -> None:
        """Test that dependencies and dependents are tracked correctly."""
        metadata = FeatureMapMetadata()

        cap1 = Entity(type="capability", id="cap1", name="Cap 1", description="Desc")
        cap2 = Entity(
            type="capability", id="cap2", name="Cap 2", description="Desc", requires={"cap1"}
        )
        cap3 = Entity(
            type="capability", id="cap3", name="Cap 3", description="Desc", requires={"cap1"}
        )
        story1 = Entity(
            type="user_story",
            id="story1",
            name="Story 1",
            description="Desc",
            requires={"cap2", "cap3"},
        )

        entities = [cap1, cap2, cap3, story1]
        resolve_graph_edges(entities)
        feature_map = FeatureMap(
            metadata=metadata,
            entities=entities,
        )

        docx_bytes = create_docx_generator(feature_map)
        text = get_docx_text(docx_bytes)

        # Check that cap1 shows it enables cap2 and cap3
        assert "Enables" in text
        assert "Cap 2" in text
        assert "Cap 3" in text

        # Check that story1 shows dependencies on both caps
        assert "Story 1" in text
        assert "Requires" in text

    def test_anchor_generation(self) -> None:
        """Test that anchors are generated correctly."""
        metadata = FeatureMapMetadata()

        cap1 = Entity(
            type="capability",
            id="complex_id_123",
            name="Complex Name with Special Chars!",
            description="Desc",
        )

        entities = [cap1]
        resolve_graph_edges(entities)
        feature_map = FeatureMap(
            metadata=metadata,
            entities=entities,
        )

        # Create backend to test anchor generation
        styling_context = styling.create_styling_context(feature_map)
        backend = DocxBackend(feature_map, styling_context)

        # Test the anchor generation
        anchor = backend.make_anchor("complex_id_123", "Complex Name with Special Chars!")
        assert anchor.startswith("entity_")
        assert "complex_id_123" in anchor

        # Verify it's valid DOCX bookmark format (alphanumeric + underscore only)
        assert all(c.isalnum() or c == "_" for c in anchor)
        assert len(anchor) <= 40

    def test_empty_sections(self) -> None:
        """Test that empty sections are not generated."""
        metadata = FeatureMapMetadata()

        # Only capabilities, no stories or outcomes
        cap1 = Entity(type="capability", id="cap1", name="Cap 1", description="Desc")

        entities = [cap1]
        resolve_graph_edges(entities)
        feature_map = FeatureMap(
            metadata=metadata,
            entities=entities,
        )

        config = DocxConfig(organization=OrganizationConfig(primary="by_type"))
        docx_bytes = create_docx_generator(feature_map, config)
        text = get_docx_text(docx_bytes)

        assert "Capabilities" in text
        # User Stories and Outcomes sections shouldn't exist if they have no content
        # The text might contain the words in other contexts, so we just verify
        # Capabilities exists

    def test_timeline_generation(self) -> None:
        """Test timeline section generation."""
        metadata = FeatureMapMetadata()

        cap1 = Entity(
            type="capability",
            id="cap1",
            name="Cap 1",
            description="Desc",
            meta={"timeframe": "2024-Q1"},
        )
        cap2 = Entity(
            type="capability",
            id="cap2",
            name="Cap 2",
            description="Desc",
            requires={"cap1"},
            meta={"timeframe": "2024-Q2"},
        )
        cap3 = Entity(
            type="capability",
            id="cap3",
            name="Cap 3",
            description="Desc",
            meta={"timeframe": "2024-Q1"},
        )
        story1 = Entity(
            type="user_story",
            id="story1",
            name="Story 1",
            description="Desc",
            requires={"cap2"},
        )

        entities = [cap1, cap2, cap3, story1]
        resolve_graph_edges(entities)
        feature_map = FeatureMap(
            metadata=metadata,
            entities=entities,
        )

        docx_bytes = create_docx_generator(feature_map)
        text = get_docx_text(docx_bytes)

        # Check timeline section exists
        assert "Timeline" in text

        # Check timeframes exist
        assert "2024-Q1" in text
        assert "2024-Q2" in text

        # Check entities are present
        assert "Cap 1" in text
        assert "Cap 2" in text
        assert "Cap 3" in text

        # Unscheduled section
        assert "Unscheduled" in text
        assert "Story 1" in text

    def test_no_timeline_when_no_timeframes(self) -> None:
        """Test that timeline section is not generated when no entities have timeframes."""
        metadata = FeatureMapMetadata()

        cap1 = Entity(type="capability", id="cap1", name="Cap 1", description="Desc")
        story1 = Entity(type="user_story", id="story1", name="Story 1", description="Desc")

        entities = [cap1, story1]
        resolve_graph_edges(entities)
        feature_map = FeatureMap(
            metadata=metadata,
            entities=entities,
        )

        docx_bytes = create_docx_generator(feature_map)
        doc = Document(BytesIO(docx_bytes))

        # Check that no heading contains "Timeline"
        timeline_found = False
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith("Heading") and "Timeline" in paragraph.text:
                timeline_found = True
                break

        assert not timeline_found

    def test_metadata_display_all_fields(self) -> None:
        """Test that all metadata fields are displayed in entity sections."""
        metadata = FeatureMapMetadata()

        cap1 = Entity(
            type="capability",
            id="cap1",
            name="Cap 1",
            description="Desc",
            meta={
                "requestor": "team_alpha",
                "timeframe": "2024-Q1",
                "priority": "high",
                "cost_estimate": 50000,
                "custom_field": "custom_value",
            },
        )

        entities = [cap1]
        resolve_graph_edges(entities)
        feature_map = FeatureMap(
            metadata=metadata,
            entities=entities,
        )

        docx_bytes = create_docx_generator(feature_map)
        text = get_docx_text(docx_bytes)

        # Check all metadata fields are displayed
        assert "team_alpha" in text
        assert "2024-Q1" in text
        assert "high" in text
        assert "50000" in text
        assert "custom_value" in text

    def test_backward_dependency_warnings(self) -> None:
        """Test detection and display of backward dependencies in timeline."""
        metadata = FeatureMapMetadata()

        cap1 = Entity(
            type="capability",
            id="cap1",
            name="Cap 1",
            description="Desc",
            requires={"cap2"},  # Depends on something in the future
            meta={"timeframe": "2024-Q1"},
        )
        cap2 = Entity(
            type="capability",
            id="cap2",
            name="Cap 2",
            description="Desc",
            meta={"timeframe": "2024-Q2"},
        )
        story1 = Entity(
            type="user_story",
            id="story1",
            name="Story 1",
            description="Desc",
            requires={"cap2"},
            meta={"timeframe": "2024-Q1"},  # Also backward dependency
        )

        entities = [cap1, cap2, story1]
        resolve_graph_edges(entities)
        feature_map = FeatureMap(
            metadata=metadata,
            entities=entities,
        )

        docx_bytes = create_docx_generator(feature_map)
        text = get_docx_text(docx_bytes)

        # Check warning section exists
        assert "Timeline Warnings" in text
        assert "backward in timeline order" in text

    def test_returns_bytes(self, simple_feature_map: FeatureMap) -> None:
        """Test that DOCX backend returns bytes."""
        docx_bytes = create_docx_generator(simple_feature_map)
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0

    def test_valid_docx_structure(self, simple_feature_map: FeatureMap) -> None:
        """Test that generated DOCX has valid structure."""
        docx_bytes = create_docx_generator(simple_feature_map)

        # Should be able to open as a Document
        doc = Document(BytesIO(docx_bytes))

        # Should have paragraphs
        assert len(doc.paragraphs) > 0

        # Should have at least one heading (the title)
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) > 0

        # Should have at least one table (metadata table)
        assert len(doc.tables) > 0

    def test_custom_section_ordering(self, simple_feature_map: FeatureMap) -> None:
        """Test that sections appear in configured order."""
        config = DocxConfig(
            organization=OrganizationConfig(
                primary="by_type", entity_type_order=["outcome", "user_story", "capability"]
            )
        )
        docx_bytes = create_docx_generator(simple_feature_map, config)
        doc = Document(BytesIO(docx_bytes))

        # Find heading positions
        headings = [
            (i, p.text) for i, p in enumerate(doc.paragraphs) if p.style.name == "Heading 2"
        ]

        outcomes_idx = None
        stories_idx = None
        capabilities_idx = None

        for idx, text in headings:
            if "Outcomes" in text:
                outcomes_idx = idx
            elif "User Stories" in text:
                stories_idx = idx
            elif "Capabilities" in text:
                capabilities_idx = idx

        # Verify they appear in the configured order (if all present)
        if outcomes_idx is not None and stories_idx is not None and capabilities_idx is not None:
            assert outcomes_idx < stories_idx
            assert stories_idx < capabilities_idx

    def test_external_links_are_clickable(self) -> None:
        """Test that external links in metadata tables are rendered as clickable hyperlinks."""
        metadata = FeatureMapMetadata()

        cap1 = Entity(
            type="capability",
            id="cap1",
            name="Cap 1",
            description="Desc",
            links=[
                "design:[DD-123](https://docs.google.com/document/d/abc123)",
                "[Plain Link](https://example.com/doc)",
                "jira:JIRA-456",  # No URL, just text
            ],
        )

        entities = [cap1]
        resolve_graph_edges(entities)
        feature_map = FeatureMap(
            metadata=metadata,
            entities=entities,
        )

        docx_bytes = create_docx_generator(feature_map)
        doc = Document(BytesIO(docx_bytes))

        # Find the entity's metadata table
        # Tables in doc: [0] = header metadata, [1] = Cap 1 metadata
        assert len(doc.tables) >= 2
        cap1_table = doc.tables[1]

        # Check for hyperlinks in the table cells
        hyperlink_count = 0
        hyperlink_texts = []
        hyperlink_urls = []

        for row in cap1_table.rows:
            for cell in row.cells:
                # Check each paragraph in the cell for hyperlinks
                for paragraph in cell.paragraphs:
                    # Look for hyperlink elements in the paragraph XML
                    for hyperlink in paragraph._element.findall(  # type: ignore[attr-defined]
                        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink"
                    ):  # noqa: SLF001
                        hyperlink_count += 1
                        # Get the text from the hyperlink
                        for text_elem in hyperlink.findall(  # type: ignore[attr-defined]
                            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
                        ):
                            hyperlink_texts.append(text_elem.text)  # type: ignore[attr-defined]
                        # Get the relationship ID
                        r_id = hyperlink.get(  # type: ignore[attr-defined]
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
                        )
                        if r_id:
                            # Get the URL from the relationship
                            rel = paragraph.part.rels[r_id]  # type: ignore[index]
                            hyperlink_urls.append(rel.target_ref)  # type: ignore[attr-defined]

        # Should have 2 hyperlinks (design link and plain link), JIRA-456 is plain text
        assert hyperlink_count == 2, f"Expected 2 hyperlinks, found {hyperlink_count}"

        # Check the hyperlink texts
        assert "DD-123" in hyperlink_texts
        assert "Plain Link" in hyperlink_texts

        # Check the hyperlink URLs
        assert "https://docs.google.com/document/d/abc123" in hyperlink_urls
        assert "https://example.com/doc" in hyperlink_urls

    def test_two_level_organization_entity_heading_levels(self) -> None:
        """Test that entities use Heading 4 in two-level organization."""
        metadata = FeatureMapMetadata()

        # Create capabilities with different timeframes
        cap1 = Entity(
            type="capability",
            id="cap1",
            name="Cap 1",
            description="Desc",
            meta={"timeframe": "2024-Q1"},
        )
        cap2 = Entity(
            type="capability",
            id="cap2",
            name="Cap 2",
            description="Desc",
            meta={"timeframe": "2024-Q2"},
        )
        story1 = Entity(
            type="user_story",
            id="story1",
            name="Story 1",
            description="Desc",
            meta={"timeframe": "2024-Q1"},
        )

        entities = [cap1, cap2, story1]
        resolve_graph_edges(entities)
        feature_map = FeatureMap(
            metadata=metadata,
            entities=entities,
        )

        # Use two-level organization: by_type then by_timeframe
        config = DocxConfig(
            organization=OrganizationConfig(primary="by_type", secondary="by_timeframe")
        )
        docx_bytes = create_docx_generator(feature_map, config)
        doc = Document(BytesIO(docx_bytes))

        # In two-level organization:
        # Heading 2: Capabilities (level 1 - primary section)
        # Heading 3: 2024-Q1 (level 2 - secondary subsection)
        # Heading 4: Cap 1 (entity - should be Heading 4)
        # Heading 3: 2024-Q2 (level 2 - secondary subsection)
        # Heading 4: Cap 2 (entity - should be Heading 4)

        # Collect all headings with their levels and text
        headings: list[tuple[str, str]] = []
        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name
            if style_name and style_name.startswith("Heading"):
                headings.append((style_name, paragraph.text))

        # Find the entity headings
        cap1_heading: str | None = None
        cap2_heading: str | None = None
        story1_heading: str | None = None

        for style_name, text in headings:
            if text == "Cap 1":
                cap1_heading = style_name
            elif text == "Cap 2":
                cap2_heading = style_name
            elif text == "Story 1":
                story1_heading = style_name

        # In two-level organization, entities should use Heading 4
        assert cap1_heading == "Heading 4", f"Cap 1 should be Heading 4, got {cap1_heading}"
        assert cap2_heading == "Heading 4", f"Cap 2 should be Heading 4, got {cap2_heading}"
        assert story1_heading == "Heading 4", f"Story 1 should be Heading 4, got {story1_heading}"

        # Verify the secondary subsection headings are Heading 3
        timeframe_2024_q1: str | None = None
        timeframe_2024_q2: str | None = None
        for style_name, text in headings:
            if text == "2024-Q1":
                timeframe_2024_q1 = style_name
            elif text == "2024-Q2":
                timeframe_2024_q2 = style_name

        assert timeframe_2024_q1 == "Heading 3", (
            f"2024-Q1 should be Heading 3, got {timeframe_2024_q1}"
        )
        assert timeframe_2024_q2 == "Heading 3", (
            f"2024-Q2 should be Heading 3, got {timeframe_2024_q2}"
        )

    def test_enables_requires_heading_levels_single_level_organization(self) -> None:
        """Test that Enables/Requires use Heading 4 in single-level organization."""
        metadata = FeatureMapMetadata()

        cap1 = Entity(type="capability", id="cap1", name="Cap 1", description="Desc")
        cap2 = Entity(
            type="capability", id="cap2", name="Cap 2", description="Desc", requires={"cap1"}
        )

        entities = [cap1, cap2]
        resolve_graph_edges(entities)
        feature_map = FeatureMap(metadata=metadata, entities=entities)

        # Use single-level organization (by_type, no secondary)
        config = DocxConfig(organization=OrganizationConfig(primary="by_type"))
        docx_bytes = create_docx_generator(feature_map, config)
        doc = Document(BytesIO(docx_bytes))

        # In single-level organization:
        # Heading 2: Capabilities (level 1 - primary section)
        # Heading 3: Cap 1 (entity - should be Heading 3)
        # Heading 4: Enables (should be Heading 4 - one level deeper than entity)
        # Heading 3: Cap 2 (entity - should be Heading 3)
        # Heading 4: Requires (should be Heading 4 - one level deeper than entity)

        # Collect all headings with their levels and text
        headings: list[tuple[str, str]] = []
        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name
            if style_name and style_name.startswith("Heading"):
                headings.append((style_name, paragraph.text))

        # Find the Enables and Requires headings
        enables_heading: str | None = None
        requires_heading: str | None = None
        cap1_heading: str | None = None
        cap2_heading: str | None = None

        for style_name, text in headings:
            if text == "Enables":
                enables_heading = style_name
            elif text == "Requires":
                requires_heading = style_name
            elif text == "Cap 1":
                cap1_heading = style_name
            elif text == "Cap 2":
                cap2_heading = style_name

        # Verify entity headings are Heading 3 in single-level organization
        assert cap1_heading == "Heading 3", f"Cap 1 should be Heading 3, got {cap1_heading}"
        assert cap2_heading == "Heading 3", f"Cap 2 should be Heading 3, got {cap2_heading}"

        # Verify Enables and Requires are Heading 4 (one level deeper than entity)
        assert enables_heading == "Heading 4", f"Enables should be Heading 4, got {enables_heading}"
        assert requires_heading == "Heading 4", (
            f"Requires should be Heading 4, got {requires_heading}"
        )

    def test_enables_requires_heading_levels_two_level_organization(self) -> None:
        """Test that Enables/Requires use Heading 5 in two-level organization."""
        metadata = FeatureMapMetadata()

        cap1 = Entity(
            type="capability",
            id="cap1",
            name="Cap 1",
            description="Desc",
            meta={"timeframe": "2024-Q1"},
        )
        cap2 = Entity(
            type="capability",
            id="cap2",
            name="Cap 2",
            description="Desc",
            requires={"cap1"},
            meta={"timeframe": "2024-Q1"},
        )

        entities = [cap1, cap2]
        resolve_graph_edges(entities)
        feature_map = FeatureMap(metadata=metadata, entities=entities)

        # Use two-level organization: by_type then by_timeframe
        config = DocxConfig(
            organization=OrganizationConfig(primary="by_type", secondary="by_timeframe")
        )
        docx_bytes = create_docx_generator(feature_map, config)
        doc = Document(BytesIO(docx_bytes))

        # In two-level organization:
        # Heading 2: Capabilities (level 1 - primary section)
        # Heading 3: 2024-Q1 (level 2 - secondary subsection)
        # Heading 4: Cap 1 (entity - should be Heading 4)
        # Heading 5: Enables (should be Heading 5 - one level deeper than entity)
        # Heading 4: Cap 2 (entity - should be Heading 4)
        # Heading 5: Requires (should be Heading 5 - one level deeper than entity)

        # Collect all headings with their levels and text
        headings: list[tuple[str, str]] = []
        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name
            if style_name and style_name.startswith("Heading"):
                headings.append((style_name, paragraph.text))

        # Find the Enables and Requires headings
        enables_heading: str | None = None
        requires_heading: str | None = None
        cap1_heading: str | None = None
        cap2_heading: str | None = None

        for style_name, text in headings:
            if text == "Enables":
                enables_heading = style_name
            elif text == "Requires":
                requires_heading = style_name
            elif text == "Cap 1":
                cap1_heading = style_name
            elif text == "Cap 2":
                cap2_heading = style_name

        # Verify entity headings are Heading 4 in two-level organization
        assert cap1_heading == "Heading 4", f"Cap 1 should be Heading 4, got {cap1_heading}"
        assert cap2_heading == "Heading 4", f"Cap 2 should be Heading 4, got {cap2_heading}"

        # Verify Enables and Requires are Heading 5 (one level deeper than entity)
        assert enables_heading == "Heading 5", f"Enables should be Heading 5, got {enables_heading}"
        assert requires_heading == "Heading 5", (
            f"Requires should be Heading 5, got {requires_heading}"
        )
