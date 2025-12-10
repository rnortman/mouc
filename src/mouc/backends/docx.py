"""DOCX backend for document generation."""
# pyright: reportUnknownVariableType=false, reportUnknownMemberType=false, reportPrivateUsage=false, reportUnknownArgumentType=false

from __future__ import annotations

from io import BytesIO
from typing import TYPE_CHECKING, Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from mouc import styling
from mouc.backends import text_formatter
from mouc.backends.base import EntityReference
from mouc.backends.text_formatter import (
    BlockElement,
    CodeBlock,
    InlineBold,
    InlineCode,
    InlineElement,
    InlineItalic,
    InlineLink,
    InlineText,
    List,
    ListItem,
    Paragraph,
)
from mouc.models import Link

if TYPE_CHECKING:
    from docx.document import Document as DocumentType
    from docx.text.paragraph import Paragraph as DocxParagraph

    from mouc.models import Entity, FeatureMap, FeatureMapMetadata
    from mouc.styling import StylingContext


class DocxBackend:
    """Backend for generating DOCX documents."""

    def __init__(
        self,
        feature_map: FeatureMap,
        styling_context: StylingContext,
        table_style: str = "Table Grid",
    ):
        """Initialize DOCX backend.

        Args:
            feature_map: The feature map being documented
            styling_context: Styling context for applying user customizations
            table_style: Word built-in table style name (default: "Table Grid")
        """
        self.feature_map = feature_map
        self.styling_context = styling_context
        self.table_style = table_style
        self.document: DocumentType | None = None
        self.bookmarks: dict[str, str] = {}  # anchor_id -> bookmark_name

    def create_document(self) -> None:
        """Initialize a new DOCX document."""
        self.document = Document()
        self.bookmarks = {}

    def add_header(self, metadata: FeatureMapMetadata) -> None:
        """Add document header with feature map metadata."""
        if not self.document:
            return

        # Add title
        self.document.add_heading("Feature Map", level=1)

        # Add metadata table
        table = self.document.add_table(rows=0, cols=2)
        table.style = self.table_style

        if metadata.team:
            row = table.add_row()
            row.cells[0].text = "Team"
            row.cells[1].text = metadata.team

        if metadata.last_updated:
            row = table.add_row()
            row.cells[0].text = "Last Updated"
            row.cells[1].text = metadata.last_updated

        row = table.add_row()
        row.cells[0].text = "Version"
        row.cells[1].text = metadata.version

    def add_section_header(self, text: str, level: int) -> str:
        """Add a section header at the specified level.

        Args:
            text: Section heading text
            level: Heading level (1=top-level, 2=subsection, etc.)

        Returns:
            Anchor ID for cross-referencing this section
        """
        if not self.document:
            return ""

        # Add heading (+1 because main title is level 1)
        heading = self.document.add_heading(text, level=level + 1)

        # Create anchor
        anchor_id = self._make_anchor_from_text(text)
        self._add_bookmark(heading, anchor_id)

        return anchor_id

    def add_entity(  # noqa: PLR0913 - Entity rendering requires multiple structured parameters
        self,
        entity: Entity,
        anchor_id: str,
        styled_type_label: str | None,
        display_metadata: dict[str, Any],
        requires_refs: list[EntityReference],
        enables_refs: list[EntityReference],
        level: int,
    ) -> None:
        """Render a complete entity in DOCX format."""
        if not self.document:
            return

        # level directly represents the heading level: 1=Heading 1, 2=Heading 2, 3=Heading 3, etc.
        # Examples:
        # level=3: Heading 3 entity
        # level=4: Heading 4 entity
        heading = self.document.add_heading(entity.name, level=level)
        self._add_bookmark(heading, anchor_id)

        # Build metadata table
        table = self.document.add_table(rows=0, cols=2)
        table.style = self.table_style

        # Add all metadata fields (including id, tags, links which are now in display_metadata)
        for key, value in sorted(display_metadata.items()):
            # Special handling for certain fields
            if key == "tags" and isinstance(value, list):
                # Tags are rendered as comma-separated plain text
                row = table.add_row()
                row.cells[0].text = key.replace("_", " ").title()
                row.cells[1].text = ", ".join(value)  # type: ignore[arg-type]
            elif key == "links" and isinstance(value, list):
                # Links get special hyperlink formatting via _format_links
                link_rows = self._format_links(value)  # type: ignore[arg-type]
                for label, link in link_rows:
                    row = table.add_row()
                    row.cells[0].text = label
                    # Add hyperlink to cell
                    self._add_link_to_cell(row.cells[1], link)
            else:
                # Regular metadata field
                row = table.add_row()
                pretty_key = key.replace("_", " ").title()
                row.cells[0].text = pretty_key
                if isinstance(value, list):
                    formatted_value = ", ".join(str(item) for item in value)  # type: ignore[arg-type]
                else:
                    formatted_value = str(value)
                row.cells[1].text = formatted_value

        # Add description (with markdown support)
        self._render_markdown_content(entity.description.strip())

        # Requires section
        if requires_refs:
            self._render_markdown_content("**Requires**")
            for ref in requires_refs:
                p = self.document.add_paragraph(style="List Bullet")
                self._add_entity_reference(p, ref, entity.type)

        # Enables section
        if enables_refs:
            self._render_markdown_content("**Enables**")
            for ref in enables_refs:
                p = self.document.add_paragraph(style="List Bullet")
                self._add_entity_reference(p, ref, entity.type)

    def add_toc_entry(
        self, text: str, anchor_id: str, level: int, suffix: str | None = None
    ) -> None:
        """Add a table of contents entry.

        Args:
            text: TOC entry text (link text)
            anchor_id: Anchor to link to
            level: Nesting level in TOC (0=top-level, 1=nested, etc.)
            suffix: Optional suffix to append after the link (e.g., type label)
        """
        if not self.document:
            return

        # Create list item with appropriate indentation
        style_name = "List Bullet" if level == 0 else f"List Bullet {level + 1}"
        try:
            p = self.document.add_paragraph(style=style_name)
        except KeyError:
            # If the style doesn't exist, use default list bullet
            p = self.document.add_paragraph(style="List Bullet")

        # Add hyperlink to bookmark
        self._add_hyperlink(p, anchor_id, text)

        # Add suffix if present
        if suffix:
            p.add_run(suffix)

    def add_timeline_warnings(self, warnings: list[str]) -> None:
        """Add timeline dependency warnings section."""
        if not warnings or not self.document:
            return

        self.document.add_heading("⚠️ Timeline Warnings", level=2)
        self.document.add_paragraph("The following dependencies go backward in timeline order:")

        for warning in warnings:
            self.document.add_paragraph(warning, style="List Bullet")

    def _render_markdown_content(self, text: str) -> None:
        """Render markdown-formatted text as DOCX elements.

        Args:
            text: Markdown-formatted text to render
        """
        if not self.document:
            return

        # Parse markdown into structured blocks
        blocks = text_formatter.parse_markdown(text)

        # Render each block
        for block in blocks:
            self._render_block(block)

    def _render_block(self, block: BlockElement) -> None:
        """Render a single block element.

        Args:
            block: Block element to render
        """
        if isinstance(block, Paragraph):
            self._render_paragraph(block)
        elif isinstance(block, CodeBlock):
            self._render_code_block(block)
        elif isinstance(block, List):
            self._render_list(block)
        elif isinstance(block, ListItem):  # type: ignore[reportUnnecessaryIsInstance]
            # List items are rendered by their parent List
            for child in block.children:
                self._render_block(child)

    def _render_paragraph(self, para: Paragraph) -> None:
        """Render a paragraph with inline formatting.

        Args:
            para: Paragraph element to render
        """
        if not self.document:
            return

        p = self.document.add_paragraph()
        for inline in para.children:
            self._render_inline_content(p, inline)

    def _render_inline_content(self, paragraph: DocxParagraph, inline: InlineElement) -> None:  # type: ignore[name-defined]
        """Render inline content (text, bold, italic, links, code) within a paragraph.

        Args:
            paragraph: DOCX paragraph to add content to
            inline: Inline element to render
        """
        if isinstance(inline, InlineText):
            paragraph.add_run(inline.text)
        elif isinstance(inline, InlineBold):
            for child in inline.children:
                if isinstance(child, InlineText):
                    run = paragraph.add_run(child.text)
                    run.bold = True
                else:
                    # Nested formatting - render child then apply bold
                    self._render_inline_with_format(paragraph, child, bold=True)
        elif isinstance(inline, InlineItalic):
            for child in inline.children:
                if isinstance(child, InlineText):
                    run = paragraph.add_run(child.text)
                    run.italic = True
                else:
                    # Nested formatting - render child then apply italic
                    self._render_inline_with_format(paragraph, child, italic=True)
        elif isinstance(inline, InlineCode):
            run = paragraph.add_run(inline.text)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            # Light gray background
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif isinstance(inline, InlineLink):  # type: ignore[reportUnnecessaryIsInstance]
            # Get text from children
            link_text = self._extract_text_from_inlines(inline.children)
            # Create hyperlink
            self._add_external_hyperlink(paragraph, inline.url, link_text)

    def _render_inline_with_format(
        self,
        paragraph: DocxParagraph,  # type: ignore[name-defined]
        inline: InlineElement,
        bold: bool = False,
        italic: bool = False,
    ) -> None:
        """Render inline content with specific formatting applied.

        Args:
            paragraph: DOCX paragraph to add content to
            inline: Inline element to render
            bold: Whether to apply bold formatting
            italic: Whether to apply italic formatting
        """
        if isinstance(inline, InlineText):
            run = paragraph.add_run(inline.text)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
        elif isinstance(inline, InlineBold):
            for child in inline.children:
                self._render_inline_with_format(paragraph, child, bold=True, italic=italic)
        elif isinstance(inline, InlineItalic):
            for child in inline.children:
                self._render_inline_with_format(paragraph, child, bold=bold, italic=True)
        elif isinstance(inline, InlineCode):
            run = paragraph.add_run(inline.text)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
        elif isinstance(inline, InlineLink):  # type: ignore[reportUnnecessaryIsInstance]
            link_text = self._extract_text_from_inlines(inline.children)
            # External hyperlinks don't support bold/italic easily - just create the link
            self._add_external_hyperlink(paragraph, inline.url, link_text)

    def _extract_text_from_inlines(self, inlines: list[InlineElement]) -> str:
        """Extract plain text from a list of inline elements.

        Args:
            inlines: List of inline elements

        Returns:
            Concatenated plain text
        """
        text_parts: list[str] = []
        for inline in inlines:
            if isinstance(inline, (InlineText, InlineCode)):
                text_parts.append(inline.text)
            elif isinstance(inline, (InlineBold, InlineItalic, InlineLink)):  # type: ignore[reportUnnecessaryIsInstance]
                text_parts.append(self._extract_text_from_inlines(inline.children))
        return "".join(text_parts)

    def _render_list(self, lst: List, base_level: int = 0) -> None:  # noqa: PLR0912
        """Render a list (ordered or unordered) with proper styling.

        Args:
            lst: List element to render
            base_level: Base indentation level (0 for top-level, 1 for nested, etc.)
        """
        if not self.document:
            return

        style_name = "List Number" if lst.ordered else "List Bullet"

        # Add level suffix for nested lists
        if base_level > 0:
            # Word styles: "List Bullet 2", "List Bullet 3", etc.
            style_name = f"{style_name} {base_level + 1}"

        for item in lst.items:
            # Each list item can contain paragraphs, code blocks, or nested lists
            if len(item.children) == 1 and isinstance(item.children[0], Paragraph):
                # Simple case: single paragraph in list item
                para = item.children[0]
                try:
                    p = self.document.add_paragraph(style=style_name)
                except KeyError:
                    # Style doesn't exist, fall back to base style
                    base_style = "List Number" if lst.ordered else "List Bullet"
                    p = self.document.add_paragraph(style=base_style)

                for inline in para.children:
                    self._render_inline_content(p, inline)
            else:
                # Complex case: multiple blocks in list item
                # Render first paragraph with list style, rest as normal blocks
                first_rendered = False
                for child in item.children:
                    if isinstance(child, Paragraph) and not first_rendered:
                        try:
                            p = self.document.add_paragraph(style=style_name)
                        except KeyError:
                            base_style = "List Number" if lst.ordered else "List Bullet"
                            p = self.document.add_paragraph(style=base_style)

                        for inline in child.children:
                            self._render_inline_content(p, inline)
                        first_rendered = True
                    elif isinstance(child, List):
                        # Nested list
                        self._render_list(child, base_level + 1)
                    else:
                        self._render_block(child)

    def _render_code_block(self, code: CodeBlock) -> None:
        """Render a code block with monospace font.

        Args:
            code: Code block element to render
        """
        if not self.document:
            return

        # Add paragraph with monospace font
        p = self.document.add_paragraph()
        run = p.add_run(code.code)
        run.font.name = "Courier New"
        run.font.size = Pt(10)

        # Add light gray background shading
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), "F0F0F0")  # Light gray
        p._element.get_or_add_pPr().append(shading_elm)  # noqa: SLF001

    def finalize(self) -> bytes:
        """Finalize and return the DOCX document."""
        if not self.document:
            return b""

        buffer = BytesIO()
        self.document.save(buffer)
        return buffer.getvalue()

    def make_anchor(self, entity_id: str, entity_name: str) -> str:
        """Create a DOCX-compatible bookmark name from entity.

        Args:
            entity_id: Unique entity identifier
            entity_name: Human-readable entity name

        Returns:
            DOCX bookmark name (must start with letter, alphanumeric + underscore only)
        """
        # DOCX bookmarks must:
        # - Start with a letter
        # - Contain only letters, numbers, and underscores
        # - Be 40 characters or less
        # Use entity_id as base for uniqueness
        bookmark = f"entity_{entity_id}"
        # Replace any non-alphanumeric with underscore
        bookmark = "".join(c if c.isalnum() or c == "_" else "_" for c in bookmark)
        # Ensure it starts with a letter (it should, given our prefix)
        if not bookmark[0].isalpha():
            bookmark = "e" + bookmark
        # Truncate to 40 chars
        return bookmark[:40]

    def format_link(self, link: Link) -> str:
        """Format an external link for display.

        Note: Returns display text. Actual hyperlink creation handled separately.
        """
        return link.label

    def format_internal_reference(self, ref: EntityReference) -> str:
        """Format an internal cross-reference.

        Note: Returns display text. Actual hyperlink creation handled separately.
        """
        return ref.entity_name

    def format_type_label(self, entity: Entity) -> str:
        """Format type label with styling applied.

        Args:
            entity: Entity to get type label for

        Returns:
            Formatted type label string (may be empty if styled to hide)
        """
        # Apply user styling
        user_label = styling.apply_label_styles(entity, self.styling_context)

        # If user styling returned a label (including empty string to hide), use it
        if user_label is not None:
            return f" {user_label}" if user_label else ""

        # Otherwise use default type label
        return f" [{self._pretty_type(entity.type)}]"

    def _format_links(self, links: list[str]) -> list[tuple[str, Link]]:
        """Format links for display in a table.

        Returns:
            List of (table_label, link) tuples where table_label is the row label
        """
        if not links:
            return []

        # Parse all links
        parsed_links = [Link.parse(link) for link in links]

        # Group by type for better organization
        by_type: dict[str | None, list[Link]] = {}
        for link in parsed_links:
            by_type.setdefault(link.type, []).append(link)

        rows: list[tuple[str, Link]] = []
        for link_type, type_links in sorted(by_type.items(), key=lambda x: (x[0] is None, x[0])):
            for link in type_links:
                if link_type:
                    # Prettify type name
                    pretty_type = link_type.replace("_", " ").title()
                    table_label = pretty_type
                else:
                    table_label = "Link"

                # Return table label and link object
                rows.append((table_label, link))

        return rows

    def _pretty_type(self, entity_type: str) -> str:
        """Convert entity type to pretty display name."""
        # Simple title-case conversion: "user_story" -> "User Story"
        return entity_type.replace("_", " ").title()

    def _make_anchor_from_text(self, text: str) -> str:
        """Create a valid DOCX bookmark name from any text."""
        # Convert to valid bookmark format
        bookmark = "bm_" + text.lower()
        bookmark = "".join(c if c.isalnum() or c == "_" else "_" for c in bookmark)
        # Remove multiple consecutive underscores
        while "__" in bookmark:
            bookmark = bookmark.replace("__", "_")
        # Truncate to 40 chars
        return bookmark[:40].strip("_")

    def _add_bookmark(self, paragraph: DocxParagraph, bookmark_name: str) -> None:  # type: ignore[name-defined]
        """Add a bookmark to a paragraph for cross-referencing."""
        # Create bookmark start element
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), str(len(self.bookmarks)))
        bookmark_start.set(qn("w:name"), bookmark_name)

        # Create bookmark end element
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), str(len(self.bookmarks)))

        # Insert bookmark around the paragraph
        paragraph._p.insert(0, bookmark_start)  # noqa: SLF001
        paragraph._p.append(bookmark_end)  # noqa: SLF001

        self.bookmarks[bookmark_name] = bookmark_name

    def _add_hyperlink(self, paragraph: DocxParagraph, bookmark_name: str, text: str) -> None:  # type: ignore[name-defined]
        """Add an internal hyperlink to a bookmark."""
        # Create hyperlink element
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), bookmark_name)

        # Create run with the text
        run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")

        # Style as hyperlink (blue, underlined)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")  # Word default hyperlink blue
        r_pr.append(color)

        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)

        run.append(r_pr)

        # Add text
        text_elem = OxmlElement("w:t")
        text_elem.text = text
        run.append(text_elem)

        hyperlink.append(run)
        paragraph._p.append(hyperlink)  # noqa: SLF001

    def _add_entity_reference(
        self,
        paragraph: DocxParagraph,
        ref: EntityReference,
        current_entity_type: str,  # type: ignore[name-defined]
    ) -> None:
        """Add an entity reference with hyperlink to a paragraph."""
        # Only create link if anchor exists (filtered/missing refs have empty anchor)
        if ref.anchor_id:
            self._add_hyperlink(paragraph, ref.anchor_id, ref.entity_name)
        else:
            paragraph.add_run(ref.entity_name)

        # Add ID in parentheses
        paragraph.add_run(f" ({ref.entity_id})")

        # Add type label if different from current type
        if ref.entity_type != current_entity_type:
            type_label = self._pretty_type(ref.entity_type)
            paragraph.add_run(f" [{type_label}]")

    def _add_link_to_cell(self, cell: Any, link: Link) -> None:
        """Add a link (with optional hyperlink) to a table cell."""
        # Get the first paragraph in the cell (or create one)
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

        # If there's a URL, create a clickable hyperlink
        if link.url:
            self._add_external_hyperlink(p, link.url, link.label)
        else:
            # Just add the label as plain text
            p.add_run(link.label)

    def _add_external_hyperlink(self, paragraph: DocxParagraph, url: str, text: str) -> None:  # type: ignore[name-defined]
        """Add an external hyperlink to a paragraph.

        Creates a proper Word hyperlink with blue underlined text.
        """
        # Get or create relationship ID for the external link
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )  # type: ignore[attr-defined]

        # Create hyperlink element
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        # Create run with the text
        run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")

        # Style as hyperlink (blue, underlined) - use explicit formatting
        # Add blue color
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")  # Word default hyperlink blue
        r_pr.append(color)

        # Add underline
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)

        run.append(r_pr)

        # Add text
        text_elem = OxmlElement("w:t")
        text_elem.text = text
        run.append(text_elem)

        hyperlink.append(run)
        paragraph._p.append(hyperlink)  # noqa: SLF001
